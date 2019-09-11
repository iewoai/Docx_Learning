from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn

q_type = '10. 单选题'
q_stem = 'Notes to financial statements are beneficial in meeting the disclosure requirements of financial reporting. The notes should not be used to'
q_opt = 'A. Correct an improper presentation in the financial statements.\nB. Describe principles and methods peculiar to the industry in which the company operates, when these principles and methods are predominantly followed in that industry.\nC. Describe depreciation methods employed by the company.\nD. Describe significant accounting policies.'
q_right = '正确答案为：A'
q_tit = u'答案解析：考点：the function of footnote to financial statement翻译：会计报表附注是会计报表的重要组成部分，附注不能用于A、对财务报表列示的项目进行修正B、描述公司所处的行业情况，和在这种环境下所采用的会计原则和方法C、描述公司所采用的折旧方法D、描述公司的主要会计政策解题思路：会计报表附注是会计报表的重要组成部分，是对会计报表本身无法或难以充分表达的内容和项目所作的补充说明和详细解释。所以附注无法纠正财务报表中的错误。A选项 附注不能对的财务报表错误进行修正，故答案是A选项，B、C、D选项 描述的是财务报表的会计原则、方法，公司所处的行业情况等都是属于附注的内容之一。'
img_path = 'F:\\py学习\\selenium\\gdCMA\\329521_stem-0.png'
# 在已有的文档写入
# document = Document('test1.docx')

# 表示重新写入
document = Document()
# 全局调整正文字体样式、大小(全局调整导致图片格式出错)
document.styles['Normal'].font.name='宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
document.styles['Normal'].font.size = Pt(10.5)
# document.styles['Normal'].paragraph_format.line_spacing = Pt(22)
# for i in range(1,38):
# 	document.add_paragraph(str(i))
p = document.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('必做题')
run.font.size = Pt(16)

document.add_paragraph(q_type).paragraph_format.line_spacing = Pt(22)

document.add_paragraph(q_stem).paragraph_format.line_spacing = Pt(22)

# 修改单段字体
# document.styles['Normal'].font.name = '宋体'
# document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# document.add_paragraph(q_tit)
# p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
# run = p.add_run()

# run.font.color.rgb = RGBColor(54, 95, 145)
# run.font.size = Pt(10.5)


# 图片居中设置
pic = document.add_picture(img_path)
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# p.add_run(q_opt)
document.add_paragraph(q_opt).paragraph_format.line_spacing = Pt(22)
# p.add_run(q_right)
document.add_paragraph(q_right).paragraph_format.line_spacing = Pt(22)
# p.add_run(q_tit)
document.add_paragraph(q_tit).paragraph_format.line_spacing = Pt(22)

document.save('test5.docx')